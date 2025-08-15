"""
File service for handling file uploads and storage

This service provides comprehensive file management capabilities for the backend,
including upload handling, text extraction, image processing, and client-specific
file organization. Frontend developers can use this service through API endpoints
to manage user files.
"""

import asyncio
import aiofiles
import base64
import hashlib
import logging
import os
from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime
import mimetypes

# Add BeautifulSoup for HTML parsing
try:
    from bs4 import BeautifulSoup
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False
    logging.warning("BeautifulSoup not available. HTML parsing will be limited.")

# Add aiohttp for URL fetching
try:
    import aiohttp
    AIOHTTP_AVAILABLE = True
except ImportError:
    AIOHTTP_AVAILABLE = False
    logging.warning("aiohttp not available. URL fetching will not be supported.")

from src.core.config import Settings
from src.models.message_models import FileInfo

# Optional PDF and DOCX parsing libraries
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDFMINER_AVAILABLE = True
except Exception:
    PDFMINER_AVAILABLE = False
    logging.warning("pdfminer.six not available. PDF text extraction will be disabled.")

try:
    import docx  # python-docx
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX text extraction will be disabled.")

logger = logging.getLogger(__name__)

class FileService:
    """
    Service for handling file operations
    
    This service manages all file-related operations including:
    - File uploads and storage with client-specific organization
    - Text extraction from various file formats (txt, md, html, pdf, docx)
    - Image extraction and processing from HTML content
    - URL content fetching and parsing
    - File cleanup and storage management
    
    Frontend Integration Notes:
    - Use client_id to associate files with specific users/sessions
    - Files are automatically organized in client-specific folders
    - Supports base64 file uploads from frontend
    - Provides file download capabilities with secure path resolution
    """
    
    def __init__(self):
        self.settings = Settings()
        # In-memory storage for client file associations
        # Key: client_id, Value: List of FileInfo objects
        self.client_files: Dict[str, List[FileInfo]] = {}
        
        # Create necessary directories for file storage
        os.makedirs(self.settings.UPLOAD_DIR, exist_ok=True)
        os.makedirs(self.settings.TEMP_DIR, exist_ok=True)
    
    async def save_file(
        self, 
        filename: str, 
        content: str, 
        file_type: str,
        client_id: Optional[str] = None
    ) -> 'FileInfo':
        """
        Save an uploaded file and return FileInfo object
        
        This method replaces save_uploaded_file to align with WebSocket message handling.
        Returns FileInfo object instead of Path for better integration.
        """
        file_path = await self.save_uploaded_file(filename, content, file_type, client_id)
        
        # Return FileInfo object for WebSocket compatibility
        return FileInfo(
            filename=filename,
            file_path=str(file_path),
            file_size=len(base64.b64decode(content)),
            file_type=file_type,
            upload_time=datetime.now().isoformat()
        )
    
    async def save_uploaded_file(
        self, 
        filename: str, 
        content: str, 
        file_type: str,
        client_id: Optional[str] = None
    ) -> Path:
        """
        Save an uploaded file to disk with client-specific organization
        
        Frontend Usage:
        - Send files as base64 encoded content
        - Always include client_id to associate files with specific users
        - Supported file types are defined in settings.ALLOWED_FILE_TYPES
        - Maximum file size is defined in settings.MAX_FILE_SIZE
        
        Args:
            filename: Original filename from frontend
            content: Base64 encoded file content
            file_type: MIME type of the file
            client_id: Unique identifier for the client/user
            
        Returns:
            Path: Path to the saved file
            
        Raises:
            ValueError: If file size exceeds limit or file type not allowed
        """
        try:
            # Decode base64 content from frontend
            file_content = base64.b64decode(content)
            
            # Validate file size against configured limits
            if len(file_content) > self.settings.MAX_FILE_SIZE:
                raise ValueError(f"File size exceeds maximum allowed size of {self.settings.MAX_FILE_SIZE} bytes")
            
            # Validate file type against allowed types
            if file_type not in self.settings.ALLOWED_FILE_TYPES:
                raise ValueError(f"File type {file_type} is not allowed")
            
            # Create client-specific folder for file organization
            if client_id:
                client_folder = Path(self.settings.UPLOAD_DIR) / f"client_{client_id}"
                client_folder.mkdir(parents=True, exist_ok=True)
                logger.info(f"Created/accessed folder for client {client_id}: {client_folder}")
            else:
                client_folder = Path(self.settings.UPLOAD_DIR)
            
            # Generate unique filename to prevent conflicts
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_hash = hashlib.md5(file_content).hexdigest()[:8]
            safe_filename = self._sanitize_filename(filename)
            unique_filename = f"{timestamp}_{file_hash}_{safe_filename}"
            
            # Create file path within client folder
            file_path = client_folder / unique_filename
            
            # Save file to disk asynchronously
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            # Create file info object for tracking
            file_info = FileInfo(
                filename=filename,
                file_path=str(file_path),
                file_size=len(file_content),
                file_type=file_type,
                upload_time=datetime.now().isoformat()
            )
            
            # Store file info for client in memory
            if client_id:
                if client_id not in self.client_files:
                    self.client_files[client_id] = []
                self.client_files[client_id].append(file_info)
                logger.info(f"Associated file {filename} with client {client_id}")
            
            logger.info(f"File saved successfully: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error saving file {filename}: {e}")
            raise
    
    def _sanitize_filename(self, filename: str) -> str:
        """
        Sanitize filename for safe storage
        
        Removes unsafe characters and limits length to prevent filesystem issues
        """
        # Remove or replace unsafe characters that could cause issues
        unsafe_chars = ['<', '>', ':', '"', '|', '?', '*', '\\', '/']
        safe_filename = filename
        for char in unsafe_chars:
            safe_filename = safe_filename.replace(char, '_')
        
        # Limit length to prevent filesystem issues
        if len(safe_filename) > 100:
            name, ext = os.path.splitext(safe_filename)
            safe_filename = name[:100-len(ext)] + ext
        
        return safe_filename
    
    async def get_client_files(self, client_id: str) -> List[FileInfo]:
        """
        Get all files uploaded by a specific client
        
        Frontend Usage:
        - Use this to display a list of uploaded files for a user
        - Returns FileInfo objects with metadata about each file
        """
        return self.client_files.get(client_id, [])
    
    async def delete_client_files(self, client_id: str) -> bool:
        """
        Delete all files for a specific client
        
        Frontend Usage:
        - Call this when a user session ends or user requests file cleanup
        - Removes both files from disk and memory
        
        Returns:
            bool: True if deletion successful, False otherwise
        """
        try:
            if client_id in self.client_files:
                # Get the client folder path
                client_folder = Path(self.settings.UPLOAD_DIR) / f"client_{client_id}"
                
                # Remove all files in the client folder
                for file_info in self.client_files[client_id]:
                    file_path = Path(file_info.file_path)
                    if file_path.exists():
                        file_path.unlink()
                        logger.info(f"Deleted file: {file_path}")
                
                # Remove the client folder if it exists and is empty
                if client_folder.exists():
                    try:
                        client_folder.rmdir()  # This will only remove if folder is empty
                        logger.info(f"Removed client folder: {client_folder}")
                    except OSError:
                        # Folder not empty, leave it for now
                        logger.info(f"Client folder not empty, leaving: {client_folder}")
                
                # Remove from memory
                del self.client_files[client_id]
                logger.info(f"Deleted all files for client {client_id}")
                return True
            return False
        except Exception as e:
            logger.error(f"Error deleting files for client {client_id}: {e}")
            return False
    
    async def get_file_content(self, file_path: str) -> Optional[bytes]:
        """
        Get file content as bytes
        
        Frontend Usage:
        - Use this to retrieve raw file content for processing or download
        """
        try:
            path = Path(file_path)
            if not path.exists():
                return None
            
            async with aiofiles.open(path, 'rb') as f:
                return await f.read()
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {e}")
            return None
    
    async def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """
        Extract text content from various file types
        
        Frontend Usage:
        - Use this to get readable text content from uploaded files
        - Supports: .txt, .md, .html, .htm, .pdf, .docx files
        
        Returns:
            str: Extracted text content or None if extraction fails
        """
        try:
            file_content = await self.get_file_content(file_path)
            if not file_content:
                return None
            
            path = Path(file_path)
            file_extension = path.suffix.lower()
            
            if file_extension == '.txt':
                return file_content.decode('utf-8', errors='ignore')
            elif file_extension == '.md':
                return file_content.decode('utf-8', errors='ignore')
            elif file_extension == '.html' or file_extension == '.htm':
                return await self._extract_text_from_html(file_content)
            elif file_extension == '.pdf':
                return await self._extract_text_from_pdf(path)
            elif file_extension == '.docx':
                return await self._extract_text_from_docx(path)
            else:
                # Try to decode as text
                return file_content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    async def extract_content_from_file(self, file_path: str) -> Optional[Dict]:
        """
        Extract both text and image content from various file types
        
        Frontend Usage:
        - Use this for comprehensive content extraction
        - Returns both text content and image metadata
        - Useful for rich content processing and display
        
        Returns:
            Dict: {
                'text': str,           # Extracted text content
                'images': List[Dict],  # List of image metadata
                'file_path': str,      # Original file path
                'file_name': str       # Original file name
            }
        """
        try:
            file_content = await self.get_file_content(file_path)
            if not file_content:
                return None
            
            path = Path(file_path)
            file_extension = path.suffix.lower()
            
            result = {
                'text': None,
                'images': [],
                'file_path': str(file_path),
                'file_name': path.name
            }
            
            if file_extension == '.txt':
                result['text'] = file_content.decode('utf-8', errors='ignore')
            elif file_extension == '.md':
                result['text'] = file_content.decode('utf-8', errors='ignore')
            elif file_extension == '.html' or file_extension == '.htm':
                result['text'] = await self._extract_text_from_html(file_content)
                result['images'] = await self._extract_images_from_html(file_content)
            elif file_extension == '.pdf':
                result['text'] = await self._extract_text_from_pdf(path)
            elif file_extension == '.docx':
                result['text'] = await self._extract_text_from_docx(path)
            else:
                # Try to decode as text
                result['text'] = file_content.decode('utf-8', errors='ignore')
            
            return result
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return None
    
    async def _extract_text_from_html(self, html_content: bytes) -> str:
        """
        Extract text content from HTML using BeautifulSoup
        
        Internal method - removes HTML tags and extracts clean text
        Requires BeautifulSoup library for full functionality
        """
        try:
            if not BEAUTIFULSOUP_AVAILABLE:
                # Fallback: try to decode as text and return raw content
                return html_content.decode('utf-8', errors='ignore')
            
            # Decode HTML content
            html_text = html_content.decode('utf-8', errors='ignore')
            
            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html_text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            # If no meaningful text found, return a fallback
            if not text.strip():
                # Try to get title
                title = soup.find('title')
                if title:
                    text = f"HTML Document: {title.get_text()}"
                else:
                    text = "HTML Document (no text content found)"
            
            return text
            
        except Exception as e:
            logger.error(f"Error parsing HTML content: {e}")
            # Fallback: return raw decoded content
            return html_content.decode('utf-8', errors='ignore')
    
    async def _extract_images_from_html(self, html_content: bytes, base_url: str = None) -> List[Dict]:
        """
        Extract image information from HTML content
        
        Internal method - finds all img tags and background images in CSS
        Returns metadata about images found in the HTML
        """
        try:
            if not BEAUTIFULSOUP_AVAILABLE:
                logger.warning("BeautifulSoup not available for image extraction")
                return []
            
            # Decode HTML content
            html_text = html_content.decode('utf-8', errors='ignore')
            
            # Parse HTML with BeautifulSoup
            soup = BeautifulSoup(html_text, 'html.parser')
            
            images = []
            
            # Find all img tags
            for img in soup.find_all('img'):
                image_info = {
                    'src': img.get('src', ''),
                    'alt': img.get('alt', ''),
                    'title': img.get('title', ''),
                    'width': img.get('width', ''),
                    'height': img.get('height', ''),
                    'class': img.get('class', []),
                    'id': img.get('id', '')
                }
                
                # Resolve relative URLs if base_url is provided
                if base_url and image_info['src']:
                    if not image_info['src'].startswith(('http://', 'https://', 'data:')):
                        from urllib.parse import urljoin
                        image_info['src'] = urljoin(base_url, image_info['src'])
                
                # Only add if we have a valid src
                if image_info['src']:
                    images.append(image_info)
            
            # Also look for background images in CSS
            for element in soup.find_all(style=True):
                style_content = element.get_text()
                # Simple regex to find background-image URLs
                import re
                bg_images = re.findall(r'background-image:\s*url\(["\']?([^"\')\s]+)["\']?\)', style_content)
                for bg_src in bg_images:
                    if bg_src and not bg_src.startswith('data:'):
                        image_info = {
                            'src': bg_src,
                            'alt': 'Background image',
                            'title': 'Background image from CSS',
                            'width': '',
                            'height': '',
                            'class': [],
                            'id': '',
                            'type': 'background'
                        }
                        
                        # Resolve relative URLs if base_url is provided
                        if base_url:
                            if not image_info['src'].startswith(('http://', 'https://')):
                                from urllib.parse import urljoin
                                image_info['src'] = urljoin(base_url, image_info['src'])
                        
                        images.append(image_info)
            
            logger.info(f"Extracted {len(images)} images from HTML")
            return images
            
        except Exception as e:
            logger.error(f"Error extracting images from HTML: {e}")
            return []

    async def _extract_text_from_pdf(self, path: Path) -> str:
        """
        Extract text from a PDF file using pdfminer.six.

        This prefers library-based extraction over plain bytes decoding.
        Returns a safe fallback message if the library is unavailable or fails.
        """
        try:
            if not PDFMINER_AVAILABLE:
                return f"[PDF parsing unavailable] Install pdfminer.six to enable parsing: {path.name}"

            # pdfminer works on file paths; handle large files efficiently
            text = pdf_extract_text(str(path)) or ""

            # Normalize whitespace
            normalized = "\n".join(
                line.strip() for line in text.splitlines() if line.strip()
            )

            # Provide a minimal fallback if empty
            return normalized if normalized else f"[No extractable text in PDF: {path.name}]"

        except Exception as e:
            logger.error(f"Error extracting PDF text from {path}: {e}")
            return f"[Failed to parse PDF: {path.name}]"

    async def _extract_text_from_docx(self, path: Path) -> str:
        """
        Extract text from a DOCX file using python-docx.

        Concatenates paragraph text preserving basic paragraph breaks.
        """
        try:
            if not DOCX_AVAILABLE:
                return f"[DOCX parsing unavailable] Install python-docx to enable parsing: {path.name}"

            document = docx.Document(str(path))
            paragraphs: List[str] = []

            for para in document.paragraphs:
                text = (para.text or "").strip()
                if text:
                    paragraphs.append(text)

            # Include table text if present
            for table in getattr(document, 'tables', []):
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = "\n".join(
                            (p.text or "").strip() for p in cell.paragraphs if (p.text or "").strip()
                        )
                        if cell_text:
                            row_cells.append(cell_text)
                    if row_cells:
                        paragraphs.append(" | ".join(row_cells))

            joined = "\n".join(paragraphs)
            return joined if joined.strip() else f"[No extractable text in DOCX: {path.name}]"

        except Exception as e:
            logger.error(f"Error extracting DOCX text from {path}: {e}")
            return f"[Failed to parse DOCX: {path.name}]"
    
    async def download_image(self, image_url: str, save_path: Path = None) -> Optional[Path]:
        """
        Download an image from URL
        
        Frontend Usage:
        - Use this to download images from external URLs
        - Images are saved to temp directory if no path specified
        - Requires aiohttp library for functionality
        
        Args:
            image_url: URL of the image to download
            save_path: Optional custom save path
            
        Returns:
            Path: Path to downloaded image or None if failed
        """
        try:
            if not AIOHTTP_AVAILABLE:
                logger.error("aiohttp not available for image downloading")
                return None
            
            async with aiohttp.ClientSession() as session:
                async with session.get(image_url, timeout=30) as response:
                    if response.status == 200:
                        image_content = await response.read()
                        
                        # Generate filename if not provided
                        if not save_path:
                            filename = image_url.split('/')[-1]
                            if '?' in filename:
                                filename = filename.split('?')[0]
                            if not filename or '.' not in filename:
                                filename = f"image_{hash(image_url) % 10000}.jpg"
                            
                            save_path = Path(self.settings.TEMP_DIR) / filename
                        
                        # Save image
                        async with aiofiles.open(save_path, 'wb') as f:
                            await f.write(image_content)
                        
                        logger.info(f"Downloaded image: {save_path}")
                        return save_path
                    else:
                        logger.error(f"Failed to download image {image_url}: HTTP {response.status}")
                        return None
                        
        except Exception as e:
            logger.error(f"Error downloading image {image_url}: {e}")
            return None
    
    async def fetch_and_parse_html_from_url(self, url: str) -> Optional[Dict]:
        """
        Fetch HTML content from a URL and extract text and images
        
        Frontend Usage:
        - Use this to extract content from web pages
        - Returns both text content and image metadata from the URL
        - Useful for web scraping and content analysis features
        
        Args:
            url: Web URL to fetch and parse
            
        Returns:
            Dict: {
                'text': str,           # Extracted text content
                'images': List[Dict],  # List of image metadata with resolved URLs
                'url': str,            # Original URL
                'content_length': int  # Size of HTML content
            }
        """
        try:
            if not AIOHTTP_AVAILABLE:
                logger.error("aiohttp not available for URL fetching")
                return None
            
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=30) as response:
                    if response.status == 200:
                        html_content = await response.read()
                        
                        # Extract text
                        text_content = await self._extract_text_from_html(html_content)
                        
                        # Extract images
                        images = await self._extract_images_from_html(html_content, url)
                        
                        return {
                            'text': text_content,
                            'images': images,
                            'url': url,
                            'content_length': len(html_content)
                        }
                    else:
                        logger.error(f"Failed to fetch URL {url}: HTTP {response.status}")
                        return None
                        
        except Exception as e:
            logger.error(f"Error fetching HTML from URL {url}: {e}")
            return None
    
    async def get_file_info(self, file_path: str) -> Optional[Dict]:
        """
        Get detailed information about a file
        
        Frontend Usage:
        - Use this to get file metadata for display
        - Returns file size, type, creation/modification times
        """
        try:
            path = Path(file_path)
            if not path.exists():
                return None
            
            stat = path.stat()
            return {
                "filename": path.name,
                "file_path": str(path),
                "file_size": stat.st_size,
                "file_type": mimetypes.guess_type(str(path))[0] or "application/octet-stream",
                "created_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat()
            }
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {e}")
            return None
    
    async def cleanup_temp_files(self, max_age_hours: int = 24) -> int:
        """
        Clean up temporary files older than specified age
        
        Admin/Maintenance Usage:
        - Call this periodically to clean up old temporary files
        - Helps manage disk space usage
        
        Returns:
            int: Number of files deleted
        """
        try:
            temp_dir = Path(self.settings.TEMP_DIR)
            if not temp_dir.exists():
                return 0
            
            cutoff_time = datetime.now().timestamp() - (max_age_hours * 3600)
            deleted_count = 0
            
            for file_path in temp_dir.rglob('*'):
                if file_path.is_file():
                    if file_path.stat().st_mtime < cutoff_time:
                        file_path.unlink()
                        deleted_count += 1
            
            logger.info(f"Cleaned up {deleted_count} temporary files")
            return deleted_count
            
        except Exception as e:
            logger.error(f"Error cleaning up temp files: {e}")
            return 0
    
    def get_storage_stats(self) -> Dict:
        """
        Get storage statistics
        
        Admin/Dashboard Usage:
        - Use this to display storage usage information
        - Shows upload/temp directory sizes and file counts
        - Useful for monitoring and analytics
        
        Returns:
            Dict: Storage statistics including sizes and counts
        """
        try:
            upload_dir = Path(self.settings.UPLOAD_DIR)
            temp_dir = Path(self.settings.TEMP_DIR)
            
            upload_size = sum(f.stat().st_size for f in upload_dir.rglob('*') if f.is_file()) if upload_dir.exists() else 0
            temp_size = sum(f.stat().st_size for f in temp_dir.rglob('*') if f.is_file()) if temp_dir.exists() else 0
            
            # Get client folder statistics
            client_folders = self.list_client_folders()
            client_folder_count = len(client_folders)
            total_client_files = sum(len(self.client_files.get(client_id, [])) for client_id in client_folders)
            
            return {
                "upload_directory_size": upload_size,
                "temp_directory_size": temp_size,
                "total_size": upload_size + temp_size,
                "client_files_count": len(self.client_files),
                "total_files_count": sum(len(files) for files in self.client_files.values()),
                "client_folders_count": client_folder_count,
                "client_folders": client_folders,
                "total_client_files": total_client_files
            }
        except Exception as e:
            logger.error(f"Error getting storage stats: {e}")
            return {}
    
    def get_client_folder_path(self, client_id: str) -> Path:
        """
        Get the folder path for a specific client
        
        Internal method - returns the file system path for client files
        """
        return Path(self.settings.UPLOAD_DIR) / f"client_{client_id}"
    
    def get_client_folder_size(self, client_id: str) -> int:
        """
        Get the total size of files in a client folder
        
        Frontend Usage:
        - Use this to show storage usage per user/client
        - Returns size in bytes
        """
        try:
            client_folder = self.get_client_folder_path(client_id)
            if not client_folder.exists():
                return 0
            
            total_size = 0
            for file_path in client_folder.rglob('*'):
                if file_path.is_file():
                    total_size += file_path.stat().st_size
            
            return total_size
        except Exception as e:
            logger.error(f"Error getting folder size for client {client_id}: {e}")
            return 0
    
    def list_client_folders(self) -> List[str]:
        """
        List all client folders
        
        Admin Usage:
        - Get list of all client IDs that have uploaded files
        - Useful for administration and analytics
        """
        try:
            upload_dir = Path(self.settings.UPLOAD_DIR)
            if not upload_dir.exists():
                return []
            
            client_folders = []
            for item in upload_dir.iterdir():
                if item.is_dir() and item.name.startswith('client_'):
                    client_id = item.name[7:]  # Remove 'client_' prefix
                    client_folders.append(client_id)
            
            return client_folders
        except Exception as e:
            logger.error(f"Error listing client folders: {e}")
            return []
    
    def cleanup_old_client_folders(self, max_age_hours: int = 24) -> int:
        """
        Clean up client folders older than specified age
        
        Admin/Maintenance Usage:
        - Call this periodically to clean up old client data
        - Helps manage disk space for inactive users
        
        Returns:
            int: Number of folders deleted
        """
        try:
            upload_dir = Path(self.settings.UPLOAD_DIR)
            if not upload_dir.exists():
                return 0
            
            cutoff_time = datetime.now().timestamp() - (max_age_hours * 3600)
            deleted_count = 0
            
            for client_folder in upload_dir.iterdir():
                if client_folder.is_dir() and client_folder.name.startswith('client_'):
                    # Check if folder is old (use modification time of any file in folder)
                    folder_age = 0
                    for file_path in client_folder.rglob('*'):
                        if file_path.is_file():
                            folder_age = max(folder_age, file_path.stat().st_mtime)
                    
                    if folder_age < cutoff_time:
                        # Remove folder and all contents
                        import shutil
                        shutil.rmtree(client_folder)
                        deleted_count += 1
                        logger.info(f"Cleaned up old client folder: {client_folder}")
            
            return deleted_count
        except Exception as e:
            logger.error(f"Error cleaning up old client folders: {e}")
            return 0
    
    def get_downloadable_file(self, file_path: str) -> tuple[Path, str]:
        """
        Resolve file path and return the actual file path and filename for download.
        
        Frontend Usage:
        - Use this to get secure file paths for download endpoints
        - Handles different path formats and prevents directory traversal attacks
        - Supports paths like: "uploads/client_123/file.txt", "client_123/file.txt", "output/file.txt"
        
        Args:
            file_path: The requested file path (can be uploads/client_*, client_*, or output/*)
            
        Returns:
            tuple: (resolved_file_path, filename)
            
        Raises:
            ValueError: If file path is invalid or access is denied
            FileNotFoundError: If file doesn't exist
        """
        try:
            logger.info(f"Resolving download path for: {file_path}")
            
            # Handle different path formats
            if file_path.startswith("uploads/"):
                # Remove the uploads/ prefix and handle as client folder
                relative_path = file_path[8:]  # Remove "uploads/"
                if relative_path.startswith("client_"):
                    upload_dir = Path(self.settings.UPLOAD_DIR)
                    requested_file = upload_dir / relative_path
                    
                    # Prevent directory traversal attacks
                    if not requested_file.resolve().is_relative_to(upload_dir.resolve()):
                        raise ValueError("Access denied: Invalid file path")
                else:
                    raise ValueError("Invalid file path format")
            elif file_path.startswith("client_"):
                # Direct client folder path
                upload_dir = Path(self.settings.UPLOAD_DIR)
                requested_file = upload_dir / file_path
                
                # Prevent directory traversal attacks
                if not requested_file.resolve().is_relative_to(upload_dir.resolve()):
                    raise ValueError("Access denied: Invalid file path")
            else:
                # File is in the output directory (backward compatibility)
                output_dir = Path("output")
                requested_file = output_dir / file_path
                
                # Prevent directory traversal attacks
                if not requested_file.resolve().is_relative_to(output_dir.resolve()):
                    raise ValueError("Access denied: Invalid file path")
            
            # Check if file exists
            if not requested_file.exists():
                raise FileNotFoundError(f"File not found: {requested_file}")
            
            # Get filename
            filename = requested_file.name
            
            logger.info(f"Resolved file path: {requested_file} -> {filename}")
            return requested_file, filename
            
        except (ValueError, FileNotFoundError):
            raise
        except Exception as e:
            logger.error(f"Error resolving file path {file_path}: {e}")
            raise ValueError(f"Error resolving file path: {e}")
    
    def check_file_exists(self, file_path: str) -> dict:
        """
        Check if a file exists and return file information.
        
        Frontend Usage:
        - Use this to verify file existence before attempting download
        - Returns detailed information about file status
        - Useful for error handling and user feedback
        
        Args:
            file_path: The requested file path
            
        Returns:
            dict: {
                'file_path': str,      # Resolved file path (if exists)
                'exists': bool,        # Whether file exists
                'size': int,           # File size in bytes
                'parent_exists': bool, # Whether parent directory exists
                'error': str           # Error message (if any)
            }
        """
        try:
            logger.info(f"Checking file existence for: {file_path}")
            
            # Try to resolve the file path
            try:
                resolved_file, filename = self.get_downloadable_file(file_path)
                exists = True
                size = resolved_file.stat().st_size
                parent_exists = resolved_file.parent.exists() if resolved_file.parent else False
            except (ValueError, FileNotFoundError):
                exists = False
                size = 0
                parent_exists = False
                resolved_file = None
            
            return {
                "file_path": str(resolved_file) if resolved_file else None,
                "exists": exists,
                "size": size,
                "parent_exists": parent_exists
            }
            
        except Exception as e:
            logger.error(f"Error checking file {file_path}: {e}")
            return {"exists": False, "error": str(e)} 